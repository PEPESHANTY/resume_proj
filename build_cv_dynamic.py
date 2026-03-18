"""
build_cv_dynamic.py  —  Generic CV renderer.
Reads content from a JSON data file; produces .docx + .pdf output.

Usage:
  # 2-page CV
  python renderer/build_cv_dynamic.py --data data/shantanu_master.json --mode 2page --company GENERIC

  # 1-page CV
  python renderer/build_cv_dynamic.py --data data/shantanu_master.json --mode 1page --company GOOGLE

  # Both (default)
  python renderer/build_cv_dynamic.py --data data/shantanu_master.json --company GENERIC

Output goes to:  CREATED/<NAME_SLUG>_CV_<COMPANY>_<MODE>.docx  (+ PDF)
"""

import argparse
import json
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "..", "CREATED")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─────────────────────────────────────────────
# STYLE CONSTANTS  (formatting only — no content)
# ─────────────────────────────────────────────

FONT_NAME  = "Calibri"
FONT_SIZE  = 10
LINK_COLOR = "0070C0"

TAB_RIGHT_POS = int(7.7 * 1440)   # right-aligned tab at end of text area


# ─────────────────────────────────────────────
# CORE HELPERS  (identical to build_cv.py)
# ─────────────────────────────────────────────

def set_page_margins(doc, top=0.5, bottom=0.25, left=0.4, right=0.4):
    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11.0)
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), 'FFFFFF')
    doc.element.insert(0, bg)
    settings = doc.settings.element
    disp_bg = OxmlElement('w:displayBackgroundShape')
    settings.insert(0, disp_bg)


def sp(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def r(para, text, bold=False, italic=False, size=FONT_SIZE, color=None, underline=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def tab_run(para, size=FONT_SIZE):
    t = para.add_run('\t')
    t.font.size = Pt(size)
    t.font.name = FONT_NAME
    return t


def hyperlink(para, text, url, size=FONT_SIZE, color=LINK_COLOR):
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    xr = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    for tag, attrs in [
        ('w:rFonts', {qn('w:ascii'): FONT_NAME, qn('w:hAnsi'): FONT_NAME}),
    ]:
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(k, v)
        rPr.append(el)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size * 2)))
    rPr.append(szCs)
    clr = OxmlElement('w:color')
    clr.set(qn('w:val'), color)
    rPr.append(clr)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    xr.insert(0, rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    xr.append(t)
    hl.append(xr)
    para._element.append(hl)
    return hl


def bottom_border(para, color="000000", space="4", size="6"):
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), size)
    b.set(qn('w:space'), space)
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)


def right_tab(para):
    pPr = para._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(TAB_RIGHT_POS))
    tabs.append(tab)
    pPr.append(tabs)


def heading(doc, title, before=5, after=2):
    p = doc.add_paragraph()
    sp(p, before=before, after=after)
    bottom_border(p)
    r(p, title, bold=True, size=FONT_SIZE)
    return p


def bul(doc, text, size=FONT_SIZE):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent       = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(0)
    r(p, text, size=size)
    return p


# ─────────────────────────────────────────────
# SECTION BUILDERS  (identical to build_cv.py)
# ─────────────────────────────────────────────

def build_header(doc, personal, name_size=18, tag_size=10.5, contact_size=9):
    p_name = doc.add_paragraph()
    sp(p_name, before=0, after=1)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p_name, personal["name"], bold=True, size=name_size)

    p_tag = doc.add_paragraph()
    sp(p_tag, before=0, after=1)
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p_tag, personal["title"], size=tag_size)

    p_c = doc.add_paragraph()
    sp(p_c, before=1, after=3)
    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep = "  |  "
    r(p_c, personal["location"] + sep, size=contact_size)
    hyperlink(p_c, personal["email"], "mailto:" + personal["email"], size=contact_size)
    r(p_c, sep + personal["phone"] + sep, size=contact_size)
    hyperlink(p_c, personal["linkedin"]["display"], personal["linkedin"]["url"], size=contact_size)
    r(p_c, sep, size=contact_size)
    hyperlink(p_c, personal["github"]["display"], personal["github"]["url"], size=contact_size)


def build_skills(doc, rows, label_width_in=2.1, before=4):
    heading(doc, "TECHNICAL SKILLS", before=before, after=2)
    content_width_in = 7.7 - label_width_in
    tbl = doc.add_table(rows=len(rows), cols=2)

    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for bname in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'bottom', 'left', 'right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tblCellMar.append(m)
    tblPr.append(tblCellMar)

    for i, (label, content) in enumerate(rows):
        row = tbl.rows[i]
        for j, w_in in enumerate([label_width_in, content_width_in]):
            tc = row.cells[j]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(w_in * 1440)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        p0 = row.cells[0].paragraphs[0]
        sp(p0, before=0, after=0)
        r(p0, label + ":", bold=True, size=FONT_SIZE)
        p1 = row.cells[1].paragraphs[0]
        sp(p1, before=0, after=0)
        r(p1, content, size=FONT_SIZE)

    p_gap = doc.add_paragraph()
    sp(p_gap, before=0, after=2)


def build_certifications(doc, certs, sec_heading="COURSES AND CERTIFICATIONS", before=4):
    heading(doc, sec_heading, before=before, after=2)
    for cert in certs:
        name      = cert["name"]
        link_text = cert["link_text"]
        url       = cert.get("url")
        date      = cert["date"]
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent       = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_before      = Pt(0)
        p.paragraph_format.space_after       = Pt(0)
        right_tab(p)
        r(p, name + " (", size=FONT_SIZE)
        if url:
            hyperlink(p, link_text, url, size=FONT_SIZE)
        else:
            r(p, link_text, size=FONT_SIZE)
        r(p, ")", size=FONT_SIZE)
        tab_run(p)
        r(p, date, bold=True, size=FONT_SIZE)


def build_education(doc, entries, before=4):
    heading(doc, "EDUCATION", before=before, after=2)
    for e in entries:
        p1 = doc.add_paragraph()
        sp(p1, before=3, after=0)
        right_tab(p1)
        r(p1, e["degree"], bold=True, size=FONT_SIZE)
        tab_run(p1)
        r(p1, e["date"], bold=True, size=FONT_SIZE)

        p2 = doc.add_paragraph()
        sp(p2, before=0, after=0)
        right_tab(p2)
        r(p2, e["institution"] + " (" + e["honors"] + ")", size=FONT_SIZE)
        tab_run(p2)
        r(p2, e["grade"], bold=True, size=FONT_SIZE)

        if e.get("modules"):
            p3 = doc.add_paragraph()
            sp(p3, before=0, after=2)
            p3.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
            label_w = Inches(1.25)
            p3.paragraph_format.left_indent       = label_w
            p3.paragraph_format.first_line_indent = -label_w
            r(p3, "Relevant modules: ", bold=True, size=FONT_SIZE)
            parts = e["modules"].split("\n")
            for i, part in enumerate(parts):
                if i > 0:
                    p3.add_run().add_break()
                r(p3, part, size=FONT_SIZE)


def build_exp_entry(doc, company, location, role, date_range, bullets):
    p_co = doc.add_paragraph()
    sp(p_co, before=4, after=0)
    right_tab(p_co)
    r(p_co, company + "  |  " + location, bold=True, size=FONT_SIZE)
    tab_run(p_co)
    r(p_co, date_range, bold=True, size=FONT_SIZE)

    p_role = doc.add_paragraph()
    sp(p_role, before=0, after=1)
    r(p_role, "Role: " + role, italic=True, size=FONT_SIZE)

    for b in bullets:
        bul(doc, b)


def build_project(doc, proj, short=False):
    p_title = doc.add_paragraph()
    sp(p_title, before=4, after=0)
    r(p_title, proj["title"], bold=True, size=FONT_SIZE)

    if proj.get("date"):
        r(p_title, "  |  ", size=FONT_SIZE)
        date_str  = proj["date"]
        paper_url = proj.get("paper_url") or ""
        link_word = "Research Paper"
        if paper_url and link_word in date_str:
            idx = date_str.index(link_word)
            before_text = date_str[:idx]
            after_text  = date_str[idx + len(link_word):]
            if before_text:
                r(p_title, before_text, bold=True, size=FONT_SIZE)
            hyperlink(p_title, link_word, paper_url, size=FONT_SIZE)
            if after_text:
                r(p_title, after_text, bold=True, size=FONT_SIZE)
        else:
            r(p_title, date_str, bold=True, size=FONT_SIZE)

    tech_str = proj.get("tech_short", proj["tech"]) if short else proj["tech"]
    p_tech = doc.add_paragraph()
    sp(p_tech, before=0, after=1)
    r(p_tech, tech_str if short else proj["tech"], italic=True, size=FONT_SIZE)

    if short:
        bullet_text = proj.get("bullet_short") or proj["bullets"][0]
        p_b = doc.add_paragraph()
        p_b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sp(p_b, before=0, after=1)
        r(p_b, bullet_text)
    else:
        for b in proj["bullets"]:
            bul(doc, b)


# ─────────────────────────────────────────────
# DOCUMENT BUILDERS
# ─────────────────────────────────────────────

def build_2page(data, output_path):
    doc = Document()
    set_page_margins(doc, top=0.5, bottom=0.25, left=0.4, right=0.4)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal'].font.size = Pt(FONT_SIZE)

    build_header(doc, data["personal"])

    if data.get("summary"):
        heading(doc, "PROFESSIONAL SUMMARY", before=4, after=2)
        p_sum = doc.add_paragraph()
        p_sum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sp(p_sum, before=0, after=2)
        r(p_sum, data["summary"])

    skill_rows = [(s["category"], s["items"]) for s in data["skills"]["full"]]
    build_skills(doc, skill_rows, label_width_in=2.1, before=4)

    build_education(doc, data["education"], before=4)

    build_certifications(doc, data["certifications"],
                         sec_heading="COURSES AND CERTIFICATIONS", before=4)

    heading(doc, "WORK EXPERIENCE", before=5, after=2)
    for exp in data["experience"]:
        build_exp_entry(doc, exp["company"], exp["location"],
                        exp["role"], exp["date_range"], exp["bullets"])

    heading(doc, "RELEVANT PROJECTS", before=5, after=2)
    for proj in data["projects"]:
        build_project(doc, proj, short=False)

    extra = data.get("extracurricular", [])
    if extra:
        heading(doc, "EXTRACURRICULAR", before=5, after=2)
        for e in extra:
            bul(doc, e)

    doc.save(output_path)
    print(f"Saved: {output_path}")


def build_1page(data, output_path, certs_count=4):
    doc = Document()
    set_page_margins(doc, top=0.45, bottom=0.2, left=0.4, right=0.4)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal'].font.size = Pt(FONT_SIZE)

    build_header(doc, data["personal"],
                 name_size=17, tag_size=10, contact_size=8.5)

    # No summary on 1-page

    # Education: strip modules for 1-page
    edu_1page = [{**e, "modules": None} for e in data["education"]]

    skill_rows = [(s["category"], s["items"]) for s in data["skills"]["compact"]]
    build_skills(doc, skill_rows, label_width_in=1.65, before=3)

    build_education(doc, edu_1page, before=3)

    heading(doc, "WORK EXPERIENCE", before=3, after=2)
    for exp in data["experience"]:
        build_exp_entry(doc, exp["company"], exp["location"],
                        exp["role"], exp["date_range"], exp["bullets_short"])

    heading(doc, "KEY PROJECTS", before=3, after=2)
    for proj in data["projects"]:
        build_project(doc, proj, short=True)

    build_certifications(doc, data["certifications"][:certs_count],
                         sec_heading="CERTIFICATIONS", before=3)

    doc.save(output_path)
    print(f"Saved: {output_path}")


# ─────────────────────────────────────────────
# PDF EXPORT
# ─────────────────────────────────────────────

def export_pdfs(docx_paths):
    import win32com.client, shutil, tempfile
    import subprocess
    subprocess.run(["taskkill", "/F", "/IM", "WINWORD.EXE"],
                   capture_output=True)
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    for docx_path in docx_paths:
        docx_path = os.path.abspath(docx_path)
        pdf_name  = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        tmp_pdf   = os.path.join(tempfile.gettempdir(), pdf_name)
        final_pdf = os.path.join(os.path.dirname(docx_path), pdf_name)
        doc = word.Documents.Open(FileName=docx_path)
        doc.ExportAsFixedFormat(tmp_pdf, ExportFormat=17)
        doc.Close(False)
        if os.path.exists(final_pdf):
            os.remove(final_pdf)
        shutil.copy2(tmp_pdf, final_pdf)
        os.remove(tmp_pdf)
        print(f"PDF: {final_pdf}")
    word.Quit()


# ─────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────

def name_slug(name):
    """'SHANTANU BHUTE' → 'SHANTANU_BHUTE'"""
    return re.sub(r'\s+', '_', name.strip().upper())


def main():
    parser = argparse.ArgumentParser(description="Build CV from JSON data file.")
    parser.add_argument("--data",    required=True,  help="Path to JSON data file")
    parser.add_argument("--company", default="GENERIC", help="Company short name for output filename")
    parser.add_argument("--mode",    default="both",  choices=["1page", "2page", "both"])
    parser.add_argument("--no-pdf",  action="store_true", help="Skip PDF export")
    args = parser.parse_args()

    with open(args.data, encoding="utf-8") as f:
        data = json.load(f)

    slug    = name_slug(data["personal"]["name"])
    company = args.company.upper()
    outputs = []

    if args.mode in ("2page", "both"):
        path = os.path.join(OUTPUT_DIR, f"{slug}_CV_{company}_2PAGE.docx")
        build_2page(data, path)
        outputs.append(path)

    if args.mode in ("1page", "both"):
        path = os.path.join(OUTPUT_DIR, f"{slug}_CV_{company}_1PAGE.docx")
        build_1page(data, path)
        outputs.append(path)

    if not args.no_pdf:
        print("Building PDFs...")
        export_pdfs(outputs)

    print("Done.")


if __name__ == "__main__":
    main()
