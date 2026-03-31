"""
build_cv_dynamic.py  —  Generic CV renderer (extended with render_config support).

This is a copy of the original renderer, extended to accept render_config.json
for per-user formatting customization. The original build_cv_dynamic.py in the
project root remains untouched.

Usage:
  # CLI mode (backward compatible)
  python backend/renderer/build_cv_dynamic.py --data data.json --mode 2page --company GOOGLE

  # With custom render config
  python backend/renderer/build_cv_dynamic.py --data data.json --config render_config.json --mode 2page

  # Programmatic usage
  from backend.renderer.build_cv_dynamic import render_cv
  paths = render_cv(data_dict, mode="2page", company="GOOGLE", config=config_dict)
"""

import argparse
import json
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ─────────────────────────────────────────────
# DEFAULT STYLE CONSTANTS  (overridden by render_config)
# ─────────────────────────────────────────────

DEFAULT_CONFIG = {
    "font_name": "Calibri",
    "font_size": 10,
    "link_color": "0070C0",
    "name_size": 18,
    "tag_size": 10.5,
    "contact_size": 9,
    "page_margins": {"top": 0.5, "bottom": 0.25, "left": 0.4, "right": 0.4},
    "section_order": ["summary", "skills", "education", "certifications", "experience", "projects", "extracurricular"],
    "include_summary": True,
    "include_extracurricular": True,
    "active": True,
    "skills_label_width": 2.1,
    "bullet_indent": 0.2,
    "max_certs_1page": 4,
    "max_projects_1page": 3,
    "page_size": "letter",
}


def _merge_config(user_config: dict | None) -> dict:
    """Merge user config on top of defaults."""
    cfg = DEFAULT_CONFIG.copy()
    if user_config:
        for k, v in user_config.items():
            if isinstance(v, dict) and isinstance(cfg.get(k), dict):
                cfg[k] = {**cfg[k], **v}
            else:
                cfg[k] = v
    return cfg


# ─────────────────────────────────────────────
# CORE HELPERS
# ─────────────────────────────────────────────

def set_page_margins(doc, cfg):
    m = cfg["page_margins"]
    pw, ph = (8.5, 11.0) if cfg["page_size"] == "letter" else (8.27, 11.69)
    for section in doc.sections:
        section.page_width    = Inches(pw)
        section.page_height   = Inches(ph)
        section.top_margin    = Inches(m["top"])
        section.bottom_margin = Inches(m["bottom"])
        section.left_margin   = Inches(m["left"])
        section.right_margin  = Inches(m["right"])
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), 'FFFFFF')
    doc.element.insert(0, bg)
    settings = doc.settings.element
    disp_bg = OxmlElement('w:displayBackgroundShape')
    settings.insert(0, disp_bg)


def sp(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def r(para, text, bold=False, italic=False, size=10, color=None, underline=False, font_name="Calibri"):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def tab_run(para, size=10, font_name="Calibri"):
    t = para.add_run('\t')
    t.font.size = Pt(size)
    t.font.name = font_name
    return t


def hyperlink(para, text, url, size=10, color="0070C0", font_name="Calibri"):
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    xr = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    for tag, attrs in [
        ('w:rFonts', {qn('w:ascii'): font_name, qn('w:hAnsi'): font_name}),
    ]:
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(k, v)
        rPr.append(el)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size * 2)))
    rPr.append(szCs)
    clr = OxmlElement('w:color')
    clr.set(qn('w:val'), color)
    rPr.append(clr)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    xr.insert(0, rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    xr.append(t)
    hl.append(xr)
    para._element.append(hl)
    return hl


def bottom_border(para, color="000000", space="4", size="6"):
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), size)
    b.set(qn('w:space'), space)
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)


def right_tab(para):
    pw = 7.7  # text area width in inches (adjustable if margins change)
    pPr = para._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(pw * 1440)))
    tabs.append(tab)
    pPr.append(tabs)


def keep_with_next(para):
    """Set 'Keep with next' on a paragraph so it won't be stranded at page bottom."""
    pPr = para._element.get_or_add_pPr()
    kwn = OxmlElement('w:keepNext')
    pPr.append(kwn)


def keep_together(para):
    """Set 'Keep lines together' so a paragraph is never split across pages."""
    pPr = para._element.get_or_add_pPr()
    klt = OxmlElement('w:keepLines')
    pPr.append(klt)


def heading_el(doc, title, cfg, before=5, after=2):
    p = doc.add_paragraph()
    sp(p, before=before, after=after)
    bottom_border(p)
    r(p, title, bold=True, size=cfg["font_size"], font_name=cfg["font_name"])
    keep_with_next(p)   # heading always stays with the content below it
    return p


def bul(doc, text, cfg):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent       = Inches(cfg["bullet_indent"])
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(0)
    r(p, text, size=cfg["font_size"], font_name=cfg["font_name"])
    return p


# ─────────────────────────────────────────────
# SECTION BUILDERS
# ─────────────────────────────────────────────

def build_header(doc, personal, cfg, is_1page=False):
    ns = cfg["name_size"] if not is_1page else cfg["name_size"] - 1
    ts = cfg["tag_size"] if not is_1page else cfg["tag_size"] - 0.5
    cs = cfg["contact_size"] if not is_1page else cfg["contact_size"] - 0.5

    p_name = doc.add_paragraph()
    sp(p_name, before=0, after=1)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p_name, personal["name"], bold=True, size=ns, font_name=cfg["font_name"])

    p_tag = doc.add_paragraph()
    sp(p_tag, before=0, after=1)
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p_tag, personal["title"], size=ts, font_name=cfg["font_name"])

    p_c = doc.add_paragraph()
    sp(p_c, before=1, after=3)
    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep = "  |  "

    r(p_c, personal["location"] + sep, size=cs, font_name=cfg["font_name"])
    hyperlink(p_c, personal["email"], "mailto:" + personal["email"],
              size=cs, color=cfg["link_color"], font_name=cfg["font_name"])
    r(p_c, sep + personal["phone"] + sep, size=cs, font_name=cfg["font_name"])

    # LinkedIn
    linkedin = personal.get("linkedin") or personal.get("links", {}).get("linkedin", {})
    if linkedin and linkedin.get("url"):
        hyperlink(p_c, linkedin.get("display", "LinkedIn"), linkedin["url"],
                  size=cs, color=cfg["link_color"], font_name=cfg["font_name"])

    # GitHub
    github = personal.get("github") or personal.get("links", {}).get("github", {})
    if github and github.get("url"):
        r(p_c, sep, size=cs, font_name=cfg["font_name"])
        hyperlink(p_c, github.get("display", "GitHub"), github["url"],
                  size=cs, color=cfg["link_color"], font_name=cfg["font_name"])


def build_summary(doc, summary_text, cfg):
    if not summary_text or not cfg.get("include_summary", True):
        return

    # Add dynamic update handling
    heading_el(doc, "PROFESSIONAL SUMMARY", cfg, before=4, after=2)
    p_sum = doc.add_paragraph()
    p_sum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sp(p_sum, before=0, after=2)
    r(p_sum, summary_text, size=cfg["font_size"], font_name=cfg["font_name"])

    # Ensure `active` filtering is respected
    if not cfg.get("active", True):
        return


def build_skills(doc, rows, cfg, is_1page=False):
    label_w = cfg.get("skills_label_width", 2.1)
    if is_1page:
        label_w = min(label_w, 1.65)
    heading_el(doc, "TECHNICAL SKILLS", cfg, before=4 if not is_1page else 3, after=2)

    content_w = 7.7 - label_w
    tbl = doc.add_table(rows=len(rows), cols=2)

    tbl_xml = tbl._tbl
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for bname in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'bottom', 'left', 'right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tblCellMar.append(m)
    tblPr.append(tblCellMar)

    for i, (label, content) in enumerate(rows):
        row = tbl.rows[i]
        for j, w_in in enumerate([label_w, content_w]):
            tc = row.cells[j]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(w_in * 1440)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        p0 = row.cells[0].paragraphs[0]
        sp(p0, before=0, after=0)
        r(p0, label + ":", bold=True, size=cfg["font_size"], font_name=cfg["font_name"])
        p1 = row.cells[1].paragraphs[0]
        sp(p1, before=0, after=0)
        r(p1, content, size=cfg["font_size"], font_name=cfg["font_name"])

    p_gap = doc.add_paragraph()
    sp(p_gap, before=0, after=2)


def build_certifications(doc, certs, cfg, sec_heading="COURSES AND CERTIFICATIONS", before=4):
    heading_el(doc, sec_heading, cfg, before=before, after=2)  # heading already has keep_with_next
    for i, cert in enumerate(certs):
        name      = cert["name"]
        link_text = cert.get("link_text", "")
        url       = cert.get("url")
        date      = cert["date"]
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent       = Inches(cfg["bullet_indent"])
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_before      = Pt(0)
        p.paragraph_format.space_after       = Pt(0)
        right_tab(p)
        # Smart link mapping: if link_text appears inside name, hyperlink that word
        if url and link_text and link_text in name:
            idx = name.index(link_text)
            before_link = name[:idx]
            after_link  = name[idx + len(link_text):]
            if before_link:
                r(p, before_link, size=cfg["font_size"], font_name=cfg["font_name"])
            hyperlink(p, link_text, url, size=cfg["font_size"],
                      color=cfg["link_color"], font_name=cfg["font_name"])
            if after_link:
                r(p, after_link, size=cfg["font_size"], font_name=cfg["font_name"])
        elif url and link_text:
            # link_text not in name — show name then (link_text) as link
            r(p, name + " (", size=cfg["font_size"], font_name=cfg["font_name"])
            hyperlink(p, link_text, url, size=cfg["font_size"],
                      color=cfg["link_color"], font_name=cfg["font_name"])
            r(p, ")", size=cfg["font_size"], font_name=cfg["font_name"])
        else:
            r(p, name, size=cfg["font_size"], font_name=cfg["font_name"])
        tab_run(p, size=cfg["font_size"], font_name=cfg["font_name"])
        r(p, date, bold=True, size=cfg["font_size"], font_name=cfg["font_name"])
        # Keep first cert bullet with heading so at least 2 items on same page
        if i < 1 and len(certs) > 1:
            keep_with_next(p)


def build_education(doc, entries, cfg, before=4):
    heading_el(doc, "EDUCATION", cfg, before=before, after=2)
    for e in entries:
        p1 = doc.add_paragraph()
        sp(p1, before=3, after=0)
        right_tab(p1)
        r(p1, e["degree"], bold=True, size=cfg["font_size"], font_name=cfg["font_name"])
        tab_run(p1, size=cfg["font_size"], font_name=cfg["font_name"])
        r(p1, e["date"], bold=True, size=cfg["font_size"], font_name=cfg["font_name"])
        keep_with_next(p1)  # degree stays with institution

        p2 = doc.add_paragraph()
        sp(p2, before=0, after=0)
        right_tab(p2)
        inst_text = e["institution"]
        if e.get("honors"):
            inst_text += " (" + e["honors"] + ")"
        r(p2, inst_text, size=cfg["font_size"], font_name=cfg["font_name"])
        if e.get("grade"):
            tab_run(p2, size=cfg["font_size"], font_name=cfg["font_name"])
            r(p2, e["grade"], bold=True, size=cfg["font_size"], font_name=cfg["font_name"])

        if e.get("modules"):
            keep_with_next(p2)  # institution stays with modules
            p3 = doc.add_paragraph()
            sp(p3, before=0, after=2)
            p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r(p3, "Relevant modules: ", bold=True, size=cfg["font_size"], font_name=cfg["font_name"])
            # Join all module lines into a single comma-separated string
            parts = [part.strip() for part in e["modules"].split("\n") if part.strip()]
            modules_text = ", ".join(parts)
            r(p3, modules_text, italic=True, size=cfg["font_size"], font_name=cfg["font_name"])


def build_exp_entry(doc, exp, cfg):
    p_co = doc.add_paragraph()
    sp(p_co, before=4, after=0)
    right_tab(p_co)
    r(p_co, exp["company"] + "  |  " + exp.get("location", ""),
      bold=True, size=cfg["font_size"], font_name=cfg["font_name"])
    tab_run(p_co, size=cfg["font_size"], font_name=cfg["font_name"])
    r(p_co, exp["date_range"], bold=True, size=cfg["font_size"], font_name=cfg["font_name"])
    keep_with_next(p_co)   # company line stays with role

    p_role = doc.add_paragraph()
    sp(p_role, before=0, after=1)
    r(p_role, "Role: " + exp["role"], italic=True, size=cfg["font_size"], font_name=cfg["font_name"])
    keep_with_next(p_role)  # role stays with first bullet

    bullets = exp.get("bullets", [])
    for i, b in enumerate(bullets):
        p = bul(doc, b, cfg)
        # Keep heading+role bound with at least 2 bullets on same page
        if i < 1 and len(bullets) > 1:
            keep_with_next(p)


def build_project(doc, proj, cfg, short=False):
    p_title = doc.add_paragraph()
    sp(p_title, before=4, after=0)
    right_tab(p_title)
    r(p_title, proj["title"], bold=True, size=cfg["font_size"], font_name=cfg["font_name"])
    keep_with_next(p_title)  # project title stays with tech line + bullets

    if proj.get("date"):
        date_str  = proj["date"]
        paper_url = proj.get("paper_url") or ""
        link_word = "Research Paper"
        # Separate the date part from link text
        if paper_url and link_word in date_str:
            idx = date_str.index(link_word)
            before_text = date_str[:idx].strip().rstrip("|").strip()
            after_text  = date_str[idx + len(link_word):].strip()
            # Title line: Project  |  date part
            r(p_title, "  |  ", size=cfg["font_size"], font_name=cfg["font_name"])
            if before_text:
                r(p_title, before_text + " ", size=cfg["font_size"], font_name=cfg["font_name"])
            hyperlink(p_title, link_word, paper_url, size=cfg["font_size"],
                      color=cfg["link_color"], font_name=cfg["font_name"])
            if after_text:
                r(p_title, " " + after_text, size=cfg["font_size"], font_name=cfg["font_name"])
            tab_run(p_title, size=cfg["font_size"], font_name=cfg["font_name"])
            # Extract just the date portion (e.g., "June 2023") and put right-aligned
            import re as _re
            date_match = _re.search(r'(\w+ \d{4})', date_str)
            if date_match:
                r(p_title, date_match.group(1), bold=True, size=cfg["font_size"], font_name=cfg["font_name"])
        else:
            tab_run(p_title, size=cfg["font_size"], font_name=cfg["font_name"])
            r(p_title, date_str, bold=True, size=cfg["font_size"], font_name=cfg["font_name"])

    tech_str = proj.get("tech_short", proj.get("tech", "")) if short else proj.get("tech", "")
    if tech_str:
        p_tech = doc.add_paragraph()
        sp(p_tech, before=0, after=1)
        r(p_tech, tech_str, italic=True, size=cfg["font_size"], font_name=cfg["font_name"])
        keep_with_next(p_tech)  # tech line stays with first bullet

    if short:
        bullet_text = proj.get("bullet_short") or (proj["bullets"][0] if proj.get("bullets") else "")
        if bullet_text:
            p_b = doc.add_paragraph()
            p_b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            sp(p_b, before=0, after=1)
            r(p_b, bullet_text, size=cfg["font_size"], font_name=cfg["font_name"])
    else:
        bullets = proj.get("bullets", [])
        for i, b in enumerate(bullets):
            p = bul(doc, b, cfg)
            # Keep title+tech bound with at least 2 bullets
            if i < 1 and len(bullets) > 1:
                keep_with_next(p)


# ─────────────────────────────────────────────
# SECTION DISPATCHER (respects section_order from config)
# ─────────────────────────────────────────────

def _normalize_items(items):
    """Ensure items is always a comma-separated string."""
    if isinstance(items, list):
        return ", ".join(str(i) for i in items)
    if isinstance(items, str):
        return items
    return str(items) if items else ""


def _get_skills_rows(data, is_1page):
    """Extract skills as list of (category, items_string) tuples from either format."""
    # Try skills_pool format first (array of {category, items})
    pool = data.get("skills_pool")
    if pool and isinstance(pool, list):
        return [(s.get("category", "Other"), _normalize_items(s.get("items", ""))) for s in pool]
    # Fall back to skills dict format {full: [...], compact: [...]}
    skills = data.get("skills") or {}
    if isinstance(skills, dict):
        source = skills.get("compact", skills.get("full", [])) if is_1page else skills.get("full", skills.get("compact", []))
        return [(s.get("category", s.get("label", "Other")), _normalize_items(s.get("items", s.get("content", "")))) for s in source]
    return []


SECTION_BUILDERS = {
    "summary": lambda doc, data, cfg, is_1page: (
        build_summary(doc, data.get("summary"), cfg) if cfg.get("active", True) else None
    ),
    "skills": lambda doc, data, cfg, is_1page: (
        build_skills(
            doc,
            _get_skills_rows(data, is_1page),
            cfg,
            is_1page=is_1page,
        ) if cfg.get("active", True) else None
    ),
    "education": lambda doc, data, cfg, is_1page: (
        build_education(
            doc,
            [{**e, "modules": None} for e in data["education"]] if is_1page else data["education"],
            cfg,
            before=3 if is_1page else 4,
        ) if cfg.get("active", True) else None
    ),
    "certifications": lambda doc, data, cfg, is_1page: (
        build_certifications(
            doc,
            data["certifications"][:cfg["max_certs_1page"]] if is_1page else data["certifications"],
            cfg,
            sec_heading="CERTIFICATIONS" if is_1page else "COURSES AND CERTIFICATIONS",
            before=3 if is_1page else 4,
        ) if cfg.get("active", True) else None
    ),
    "experience": lambda doc, data, cfg, is_1page: _build_experience_section(doc, data, cfg, is_1page),
    "projects": lambda doc, data, cfg, is_1page: _build_projects_section(doc, data, cfg, is_1page),
    "extracurricular": lambda doc, data, cfg, is_1page: _build_extra_section(doc, data, cfg, is_1page),
}


def _build_experience_section(doc, data, cfg, is_1page):
    heading_el(doc, "WORK EXPERIENCE", cfg, before=3 if is_1page else 5, after=2)
    experiences = [exp for exp in data.get("experience", []) if exp.get("active", True)]
    for exp in experiences:
        if is_1page and "bullets_short" in exp:
            exp_render = {**exp, "bullets": exp["bullets_short"]}
        else:
            exp_render = exp
        build_exp_entry(doc, exp_render, cfg)


def _build_projects_section(doc, data, cfg, is_1page):
    heading_el(doc, "KEY PROJECTS" if is_1page else "RELEVANT PROJECTS", cfg,
               before=3 if is_1page else 5, after=2)
    projects = [proj for proj in data.get("projects", []) if proj.get("active", True)]
    if is_1page:
        projects = projects[:cfg.get("max_projects_1page", 3)]
    for proj in projects:
        build_project(doc, proj, cfg, short=is_1page)


def _build_extra_section(doc, data, cfg, is_1page):
    if is_1page or not cfg.get("include_extracurricular", True):
        return
    extra = data.get("extracurricular", [])
    extra = [e for e in extra if (e.get("active", True) if isinstance(e, dict) else True)]
    if not extra:
        return
    heading_el(doc, "EXTRACURRICULAR", cfg, before=5, after=2)
    for i, e in enumerate(extra):
        text = e if isinstance(e, str) else e.get("text", "")
        p = bul(doc, text, cfg)
        # Keep first bullet with heading
        if i < 1 and len(extra) > 1:
            keep_with_next(p)


# ─────────────────────────────────────────────
# DOCUMENT BUILDERS
# ─────────────────────────────────────────────

def build_cv(data: dict, output_path: str, mode: str = "2page", config: dict | None = None):
    """
    Build a single CV document.

    Args:
        data: CV content data (schema.json format).
        output_path: Where to save the .docx.
        mode: "1page" or "2page".
        config: Optional render_config overrides.
    """
    cfg = _merge_config(config)
    is_1page = mode == "1page"

    doc = Document()
    set_page_margins(doc, cfg)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].font.name = cfg["font_name"]
    doc.styles['Normal'].font.size = Pt(cfg["font_size"])

    # Header (always first)
    build_header(doc, data["personal"], cfg, is_1page=is_1page)

    # Build sections in configured order
    section_order = cfg.get("section_order", list(SECTION_BUILDERS.keys()))
    for section_name in section_order:
        builder = SECTION_BUILDERS.get(section_name)
        if builder:
            builder(doc, data, cfg, is_1page)

    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


# ─────────────────────────────────────────────
# PDF EXPORT
# ─────────────────────────────────────────────

def export_pdfs(docx_paths):
    """Export .docx to .pdf — tries LibreOffice first (cross-platform), falls back to MS Word COM on Windows."""
    import subprocess
    pdf_paths = []
    for docx_path in docx_paths:
        docx_path = os.path.abspath(docx_path)
        out_dir   = os.path.dirname(docx_path)
        pdf_name  = os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
        pdf_path  = os.path.join(out_dir, pdf_name)

        # Try LibreOffice (works on Linux/Mac/Windows if installed)
        converted = False
        for soffice_cmd in ("soffice", "libreoffice"):
            try:
                subprocess.run(
                    [soffice_cmd, "--headless", "--convert-to", "pdf",
                     "--outdir", out_dir, docx_path],
                    capture_output=True, check=True, timeout=60,
                )
                if os.path.exists(pdf_path):
                    pdf_paths.append(pdf_path)
                    print(f"PDF (LibreOffice): {pdf_path}")
                    converted = True
                    break
            except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
                continue

        if converted:
            continue

        # Fallback: MS Word COM (Windows only)
        try:
            import win32com.client, shutil, tempfile
            subprocess.run(["taskkill", "/F", "/IM", "WINWORD.EXE"], capture_output=True)
            word    = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            tmp_pdf = os.path.join(tempfile.gettempdir(), pdf_name)
            doc     = word.Documents.Open(FileName=docx_path)
            doc.ExportAsFixedFormat(tmp_pdf, ExportFormat=17)
            doc.Close(False)
            word.Quit()
            shutil.copy2(tmp_pdf, pdf_path)
            os.remove(tmp_pdf)
            pdf_paths.append(pdf_path)
            print(f"PDF (Word COM): {pdf_path}")
        except Exception as e:
            print(f"WARNING: PDF conversion failed for {docx_path}: {e}")

    return pdf_paths


# ─────────────────────────────────────────────
# HIGH-LEVEL API  (used by FastAPI / Streamlit)
# ─────────────────────────────────────────────

def render_cv(data: dict, mode: str = "2page", company: str = "GENERIC",
              config: dict | None = None, output_dir: str | None = None,
              export_pdf: bool = False) -> dict:
    """
    Render CV and return paths to generated files.

    Returns:
        {"docx": "/path/to.docx", "pdf": "/path/to.pdf" or None}
    """
    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(__file__), "..", "..", "CREATED")
    os.makedirs(output_dir, exist_ok=True)

    slug = re.sub(r'\s+', '_', data["personal"]["name"].strip().upper())
    docx_name = f"{slug}_CV_{company.upper()}_{mode.upper()}.docx"
    docx_path = os.path.join(output_dir, docx_name)

    build_cv(data, docx_path, mode=mode, config=config)

    pdf_path = None
    if export_pdf:
        pdfs = export_pdfs([docx_path])
        if pdfs:
            pdf_path = pdfs[0]

    return {"docx": docx_path, "pdf": pdf_path}


# ─────────────────────────────────────────────
# CLI (backward compatible)
# ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Build CV from JSON data file.")
    parser.add_argument("--data",    required=True,  help="Path to JSON data file")
    parser.add_argument("--config",  default=None,   help="Optional render_config.json path")
    parser.add_argument("--company", default="GENERIC", help="Company short name")
    parser.add_argument("--mode",    default="both",  choices=["1page", "2page", "both"])
    parser.add_argument("--no-pdf",  action="store_true", help="Skip PDF export")
    parser.add_argument("--outdir",  default=None,   help="Output directory")
    args = parser.parse_args()

    with open(args.data, encoding="utf-8") as f:
        data = json.load(f)

    config = None
    if args.config:
        with open(args.config, encoding="utf-8") as f:
            config = json.load(f)

    modes = ["2page", "1page"] if args.mode == "both" else [args.mode]
    for mode in modes:
        result = render_cv(data, mode=mode, company=args.company,
                           config=config, output_dir=args.outdir,
                           export_pdf=not args.no_pdf)
        print(f"  DOCX: {result['docx']}")
        if result['pdf']:
            print(f"  PDF:  {result['pdf']}")

    print("Done.")


if __name__ == "__main__":
    main()
