"""
Build Shantanu Bhute's CV as .docx files.
Creates both 2-page and 1-page versions in d:/A RESUMES/CREATED/
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os

OUTPUT_DIR = "d:/A RESUMES/CREATED"
os.makedirs(OUTPUT_DIR, exist_ok=True)

FONT_NAME   = "Calibri"
FONT_SIZE   = 10        # all body text
LINK_COLOR  = "0070C0"  # blue matching Eaton PDF links

# ─────────────────────────────────────────────
# LINKS  (stored in memory/links_and_urls.md)
# ─────────────────────────────────────────────

SIAMESE_PAPER_URL = "https://drive.google.com/file/d/1FjL5FX5ht4I_JSSQODYNlO884R90YfNc/view?usp=sharing"

LINKEDIN_URL = "https://www.linkedin.com/in/shantanubhute/"
GITHUB_URL   = "https://github.com/PEPESHANTY"
EMAIL        = "shantanubhute@gmail.com"
PHONE        = "+353 899476612"

# (cert display name, link anchor text, url, date)
CERTS_LINKED = [
    ("AWS Certified Cloud Practitioner",
     "Certified Badge",
     "https://www.credly.com/badges/e85ed162-77af-455a-b9e4-61b12626cea3/linked_in_profile",
     "Dec 2024"),
    ("Data Science Course: UNIVERSITY OF MICHIGAN",
     "Certification link",
     "https://www.coursera.org/account/accomplishments/verify/R9WBDMY72N7S",
     "Dec 2023"),
    ("1Z0-071: Oracle Database SQL Certification",
     "Certified Badge",
     "https://catalog-education.oracle.com/ords/certview/sharebadge?id=EF592A9A40985E62B928AC4BD900A4F92CECC465A9BDA5699F83A359A03486B1",
     "Aug 2024"),
    ("AWS Academy Cloud Architecting and Accreditation",
     "Certification links",
     "https://docs.google.com/document/d/1f6vrpOaq0IW1NCCpuT2LzgXvasZiekhpjVQEKVzqVzY/edit?tab=t.0#heading=h.gh8vm1wg984a",
     "May 2022"),
    ("Linux Fundamentals",
     "Certification link",
     "https://www.coursera.org/account/accomplishments/verify/ZFJQKNXQ8BSZ",
     "Mar 2022"),
]

CERTS_1PAGE = CERTS_LINKED[:4]   # drop Linux on 1-page to save space

# ─────────────────────────────────────────────
# TAB STOP POSITIONS  (twips = inches × 1440)
# ─────────────────────────────────────────────

TAB_RIGHT_POS      = int(7.7  * 1440)   # right-aligned tab at end of text area
TAB_SKILLS_FULL    = int(2.5  * 1440)   # left tab for 2-page skills (long labels)
TAB_SKILLS_COMPACT = int(1.75 * 1440)   # left tab for 1-page skills (short labels)


# ─────────────────────────────────────────────
# CORE HELPERS
# ─────────────────────────────────────────────

def set_page_margins(doc, top=0.5, bottom=0.25, left=0.4, right=0.4):
    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11.0)
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), 'FFFFFF')
    doc.element.insert(0, bg)
    settings = doc.settings.element
    disp_bg = OxmlElement('w:displayBackgroundShape')
    settings.insert(0, disp_bg)


def sp(para, before=0, after=0, line=None):
    """Set paragraph spacing."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def r(para, text, bold=False, italic=False, size=FONT_SIZE, color=None, underline=False):
    """Add a formatted run."""
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def tab_run(para, size=FONT_SIZE):
    """Add a tab character run."""
    t = para.add_run('\t')
    t.font.size = Pt(size)
    t.font.name = FONT_NAME
    return t


def hyperlink(para, text, url, size=FONT_SIZE, color=LINK_COLOR):
    """Add a clickable hyperlink run."""
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    xr = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    for tag, attrs in [
        ('w:rFonts', {qn('w:ascii'): FONT_NAME, qn('w:hAnsi'): FONT_NAME}),
    ]:
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(k, v)
        rPr.append(el)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size * 2)))
    rPr.append(szCs)
    clr = OxmlElement('w:color')
    clr.set(qn('w:val'), color)
    rPr.append(clr)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    xr.insert(0, rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    xr.append(t)
    hl.append(xr)
    para._element.append(hl)
    return hl


def bottom_border(para, color="000000", space="4", size="6"):
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), size)
    b.set(qn('w:space'), space)
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)


def right_tab(para):
    """Apply right-aligned tab stop at end of text area."""
    pPr = para._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(TAB_RIGHT_POS))
    tabs.append(tab)
    pPr.append(tabs)


def left_tab(para, pos):
    """Apply left-aligned tab stop at given twips position."""
    pPr = para._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(pos))
    tabs.append(tab)
    pPr.append(tabs)


def heading(doc, title, before=5, after=2):
    p = doc.add_paragraph()
    sp(p, before=before, after=after)
    bottom_border(p)
    r(p, title, bold=True, size=FONT_SIZE)
    return p


def bul(doc, text, size=FONT_SIZE):
    """Justified bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent       = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(0)
    r(p, text, size=size)
    return p


# ─────────────────────────────────────────────
# HEADER
# ─────────────────────────────────────────────

def build_header(doc, tagline, name_size=18, tag_size=10.5, contact_size=9):
    p_name = doc.add_paragraph()
    sp(p_name, before=0, after=1)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p_name, "SHANTANU BHUTE", bold=True, size=name_size)

    p_tag = doc.add_paragraph()
    sp(p_tag, before=0, after=1)
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r(p_tag, tagline, size=tag_size)

    p_c = doc.add_paragraph()
    sp(p_c, before=1, after=3)
    p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep = "  |  "
    r(p_c, "Dublin, Ireland" + sep, size=contact_size)
    hyperlink(p_c, EMAIL, "mailto:" + EMAIL, size=contact_size)
    r(p_c, sep + PHONE + sep, size=contact_size)
    hyperlink(p_c, "LinkedIn", LINKEDIN_URL, size=contact_size)
    r(p_c, sep, size=contact_size)
    hyperlink(p_c, "GitHub", GITHUB_URL, size=contact_size)


# ─────────────────────────────────────────────
# TECHNICAL SKILLS  (2-column borderless table for perfect alignment)
# ─────────────────────────────────────────────

def build_skills(doc, rows, label_width_in=2.1, before=4):
    """
    Renders skills as a 2-column borderless table so content always
    starts at the same x-position regardless of label length.
    label_width_in: width of the bold-label column in inches.
    """
    heading(doc, "TECHNICAL SKILLS", before=before, after=2)

    content_width_in = 7.7 - label_width_in
    tbl = doc.add_table(rows=len(rows), cols=2)

    # ── remove all table borders ──
    tbl_xml = tbl._tbl
    # get or create tblPr
    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for bname in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # ── zero cell margins ──
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ['top', 'bottom', 'left', 'right']:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tblCellMar.append(m)
    tblPr.append(tblCellMar)

    # ── fill rows ──
    for i, (label, content) in enumerate(rows):
        row = tbl.rows[i]

        # fix column widths on each cell
        for j, w_in in enumerate([label_width_in, content_width_in]):
            tc = row.cells[j]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(w_in * 1440)))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        # label cell
        p0 = row.cells[0].paragraphs[0]
        sp(p0, before=0, after=0)
        r(p0, label + ":", bold=True, size=FONT_SIZE)

        # content cell
        p1 = row.cells[1].paragraphs[0]
        sp(p1, before=0, after=0)
        r(p1, content, size=FONT_SIZE)

    # small gap after table
    p_gap = doc.add_paragraph()
    sp(p_gap, before=0, after=2)


# ─────────────────────────────────────────────
# CERTIFICATIONS  (name (linked_text) — Eaton style)
# ─────────────────────────────────────────────

def build_certifications(doc, certs, sec_heading="COURSES AND CERTIFICATIONS", before=4):
    """
    certs: list of (name, link_text, url, date)
    Renders: • name (linked_text)    [right-aligned bold] date
    """
    heading(doc, sec_heading, before=before, after=2)
    for name, link_text, url, date in certs:
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent       = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_before      = Pt(0)
        p.paragraph_format.space_after       = Pt(0)
        right_tab(p)
        r(p, name + " (", size=FONT_SIZE)
        hyperlink(p, link_text, url, size=FONT_SIZE)
        r(p, ")", size=FONT_SIZE)
        tab_run(p)
        r(p, date, bold=True, size=FONT_SIZE)


# ─────────────────────────────────────────────
# EDUCATION  (Eaton-style: 3-line per entry)
# ─────────────────────────────────────────────

def build_education(doc, entries, before=4):
    """
    entries: list of dicts with keys:
      degree, institution, honors, grade, date, modules (optional)
    """
    heading(doc, "EDUCATION", before=before, after=2)
    for e in entries:
        # Line 1: degree (bold, left) + date (right)
        p1 = doc.add_paragraph()
        sp(p1, before=3, after=0)
        right_tab(p1)
        r(p1, e["degree"], bold=True, size=FONT_SIZE)
        tab_run(p1)
        r(p1, e["date"], bold=True, size=FONT_SIZE)

        # Line 2: institution (honors) (left) + grade (right)
        p2 = doc.add_paragraph()
        sp(p2, before=0, after=0)
        right_tab(p2)
        r(p2, e["institution"] + " (" + e["honors"] + ")", size=FONT_SIZE)
        tab_run(p2)
        r(p2, e["grade"], bold=True, size=FONT_SIZE)

        # Line 3: Relevant modules — hanging indent so wrap aligns to content start
        # Use \n in modules string to force a manual line break at that point
        if e.get("modules"):
            p3 = doc.add_paragraph()
            sp(p3, before=0, after=2)
            p3.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE  # stretches every line incl. soft-return lines to full width
            # "Relevant modules: " at 10pt Calibri bold ≈ 1.25 inches
            label_w = Inches(1.25)
            p3.paragraph_format.left_indent       = label_w
            p3.paragraph_format.first_line_indent = -label_w
            r(p3, "Relevant modules: ", bold=True, size=FONT_SIZE)
            parts = e["modules"].split("\n")
            for i, part in enumerate(parts):
                if i > 0:
                    p3.add_run().add_break()   # manual line break
                r(p3, part, size=FONT_SIZE)


# ─────────────────────────────────────────────
# WORK EXPERIENCE
# ─────────────────────────────────────────────

def build_exp_entry(doc, company, location, role, date_range, bullets):
    # Company | Location  →right→  date
    p_co = doc.add_paragraph()
    sp(p_co, before=4, after=0)
    right_tab(p_co)
    r(p_co, company + "  |  " + location, bold=True, size=FONT_SIZE)
    tab_run(p_co)
    r(p_co, date_range, bold=True, size=FONT_SIZE)

    # Role (italic)
    p_role = doc.add_paragraph()
    sp(p_role, before=0, after=1)
    r(p_role, "Role: " + role, italic=True, size=FONT_SIZE)

    for b in bullets:
        bul(doc, b)


# ─────────────────────────────────────────────
# PROJECTS  (title inline with bold context — no right tab)
# ─────────────────────────────────────────────

def build_project(doc, proj):
    p_title = doc.add_paragraph()
    sp(p_title, before=4, after=0)
    r(p_title, proj["title"], bold=True, size=FONT_SIZE)

    # Context/date: inline after separator, bold
    # If paper_url is set, "Research Paper" becomes a hyperlink
    if proj.get("date"):
        r(p_title, "  |  ", size=FONT_SIZE)
        date_str = proj["date"]
        paper_url = proj.get("paper_url", "")
        link_word = "Research Paper"
        if paper_url and link_word in date_str:
            idx = date_str.index(link_word)
            before_text = date_str[:idx]
            after_text  = date_str[idx + len(link_word):]
            if before_text:
                r(p_title, before_text, bold=True, size=FONT_SIZE)
            hyperlink(p_title, link_word, paper_url, size=FONT_SIZE)
            if after_text:
                r(p_title, after_text, bold=True, size=FONT_SIZE)
        else:
            r(p_title, date_str, bold=True, size=FONT_SIZE)

    p_tech = doc.add_paragraph()
    sp(p_tech, before=0, after=1)
    r(p_tech, proj["tech"], italic=True, size=FONT_SIZE)

    for b in proj["bullets"]:
        bul(doc, b)


# ─────────────────────────────────────────────
# CONTENT DATA
# ─────────────────────────────────────────────

# Skills — Eaton-style categories (used in 2-page)
SKILLS_FULL = [
    ("Languages",                    "Python, SQL, Java, R, Bash, HTML, CSS"),
    ("Frameworks",                   "TensorFlow, PyTorch, scikit-learn, Flask, Streamlit, Spring Boot, Machine Learning"),
    ("GenAI & LLM Tools",            "OpenAI API, LangChain, Hugging Face Hub, Prompt Engineering, Statistical Modelling"),
    ("Databases",                    "PostgreSQL, MySQL, MongoDB, Redis, AWS Redshift, AWS RDS"),
    ("Cloud & MLOps",                "AWS (SageMaker, Lambda, S3, EC2, RDS), Azure ML, Docker, Kubernetes, MLflow, Edge AI"),
    ("Data Engineering & Pipelines", "Apache Airflow, Spark, Kafka, Pandas, NumPy, AWS Glue, Prompt Tuning, ETL Workflows"),
    ("Visualization & Reporting",    "Power BI, Matplotlib, Seaborn, Plotly, ggplot2"),
    ("ML & Data Science",            "Python, R, pandas, NumPy, TensorFlow, PyTorch, scikit-learn, ggplot2"),
    ("Others",                       "REST APIs, FastAPI, OpenCV, NLP, Model Deployment (API & Streamlit), Git, OpenCV"),
]

SKILLS_COMPACT = [
    ("Languages",    "Python, SQL, Java, R, Bash"),
    ("ML & AI",      "TensorFlow, PyTorch, scikit-learn, LangChain, OpenAI API, Hugging Face, NLP, Computer Vision"),
    ("Data & Cloud", "Airflow, Spark, Kafka, Pandas, AWS (SageMaker, Lambda, S3, Redshift), Azure ML, Docker, MLflow"),
    ("Visualization","Power BI, Matplotlib, Seaborn, Plotly"),
    ("Tools",        "Streamlit, Flask, FastAPI, Spring Boot, REST APIs, Git, OpenCV, PostgreSQL, MongoDB"),
]

# Education entries
EDUCATION_2PAGE = [
    {
        "degree":      "Master of Science in Computer Science",
        "institution": "University College Dublin, Ireland",
        "honors":      "First Class with Honors",
        "grade":       "(1:1)",
        "date":        "Sept 2024 – Sept 2025",
        "modules":     "Cloud Computing, Multi-Agent Systems, Distributed Systems, Big Data Programming,\n"
                       "Generative AI, RDBMS, Machine Learning, Enterprise & Innovation",
    },
    {
        "degree":      "Bachelor of Engineering in Computer Science with Honours in AI-ML",
        "institution": "Savitribai Phule Pune University, India",
        "honors":      "First Class with Honors",
        "grade":       "CGPA: 8.97/10",
        "date":        "Aug 2019 – May 2023",
        "modules":     "Object-Oriented Programming, Cloud Computing, Blockchain, AI-ML, Web Technologies",
    },
]

EDUCATION_1PAGE = [
    {
        "degree":      "Master of Science in Computer Science",
        "institution": "University College Dublin, Ireland",
        "honors":      "First Class with Honors",
        "grade":       "(1:1)",
        "date":        "Sept 2024 – Sept 2025",
        "modules":     None,
    },
    {
        "degree":      "Bachelor of Engineering in Computer Science (AI-ML)",
        "institution": "Savitribai Phule Pune University, India",
        "honors":      "First Class with Honors",
        "grade":       "CGPA: 8.97/10",
        "date":        "Aug 2019 – May 2023",
        "modules":     None,
    },
]

# Experience bullets — Generic v2 (Current Default)
CEADAR_BULLETS = [
    "Designed and deployed a Retrieval-Augmented Generation (RAG) pipeline using LangChain, OpenAI, and Qdrant to enable domain-specific knowledge retrieval from PDFs and rice/agriculture websites",
    "Built a Streamlit and React based chatbot UI with memory, citations, and multilingual (English & Vietnamese) support, collaborating closely with cross-functional teams",
    "Automated PDF/web ingestion pipelines with chunking, embeddings (OpenAI/Qwen), metadata enrichment, and Qdrant storage, and implemented evaluation workflows using RAGAS to validate retrieval accuracy",
    "Developed a computer vision model using PyTorch to classify and label multiple rice disease varieties from field images, applying CNN-based architectures and transfer learning for early crop disease detection",
]

CEADAR_SHORT = [
    "Built production-grade RAG pipeline (LangChain, OpenAI, Qdrant) for domain-specific agricultural knowledge retrieval with multilingual (English/Vietnamese) support and RAGAS-validated benchmarking",
    "Automated PDF/web ingestion pipelines (chunking, embeddings, Qdrant/Supabase storage) and built Streamlit/React chatbot UI with persistent memory and source citations",
    "Trained CNN-based computer vision model (PyTorch, transfer learning) for multi-class rice disease classification from field images",
]

LTIMINDTREE_BULLETS = [
    "Optimized Oracle SQL queries and backend pipelines, improving data retrieval performance for Citibank's reporting systems",
    "Developed Java modules for intelligent data validation and integrated backend checks with early-stage anomaly detection logic",
    "Explored prompt engineering and Python-based model testing to support internal AI-assisted decision modules",
    "Designed and validated ML pipelines with quality checks and logging on AWS (Lambda + S3), ensuring data integrity and anomaly detection in low-latency environments",
]

LTIMINDTREE_SHORT = [
    "Optimized Oracle SQL queries and backend pipelines for Citibank's reporting systems; developed Java data validation modules with anomaly detection",
    "Designed and validated ML pipelines with quality checks and logging on AWS (Lambda + S3) for low-latency anomaly detection",
]

COGNIZANT_BULLETS = [
    "Built conversational AI chatbot using Google Dialogflow, trained on CRM use cases to automate customer query prediction",
    "Developed backend services using Java and Spring Boot to integrate chatbot intelligence with CRM workflows",
    "Improved dynamic routing and response accuracy using Avaya Aura and Genesys Cloud with integrated AI triggers",
    "Designed intent classification strategies using NLP preprocessing and training phrase optimization, improving chatbot prediction accuracy by 18% across varied customer input",
]

COGNIZANT_SHORT = [
    "Built Dialogflow conversational AI chatbot integrated with CRM workflows via Java/Spring Boot and Avaya Aura/Genesys Cloud",
    "Improved chatbot prediction accuracy by 18% through NLP intent classification redesign and training phrase optimization",
]

PROJ_RICEAI = {
    "title":   "RiceAI Expert – Sustainable Farming AI Chatbot",
    "date":    "Dec 2025",
    "tech":    "Technologies: LangChain, LangGraph, Qdrant, OpenAI/Qwen Embeddings, Streamlit, Supabase, Neo4j, RAGAS, Git",
    "bullets": [
        "Designed a full-stack RAG assistant for rice farming, integrating agricultural PDFs and web content with hybrid retrieval using LangChain and Qdrant",
        "Extended to GraphRAG (LangGraph + Neo4j) for explainable, knowledge-graph-enhanced retrieval with source traceability",
        "Delivered multilingual support (English & Vietnamese), automated ingestion pipelines, and RAGAS-validated quality benchmarking",
    ]
}

PROJ_SCHOLARGENIE = {
    "title":   "ScholarGenie – LLM-Powered Research Paper Discovery Agent",
    "date":    "March 2024",
    "tech":    "Technologies: Python, LangChain, OpenAI GPT-4, arXiv API, Semantic Scholar API, Streamlit",
    "bullets": [
        "Built an LLM-powered academic discovery tool integrating arXiv and Semantic Scholar APIs for real-time paper search, keyword filtering, and GPT-4-generated plain-language summaries",
        "Deployed with Streamlit for interactive exploration of AI-generated research summaries and trend analysis",
    ]
}

PROJ_SIAMESE = {
    "title":     "Face Verification using Siamese Convolutional Neural Network",
    "date":      "Research Paper, June 2023",
    "paper_url": SIAMESE_PAPER_URL,   # fill SIAMESE_PAPER_URL at top of file
    "tech":    "Technologies: Deep Learning, TensorFlow, Keras, Python, NumPy, OpenCV",
    "bullets": [
        "Developed a Siamese CNN for one-shot face verification, achieving 91.25% accuracy classifying 13,000 unique facial images; outperformed KNN and random-guess baselines",
        "Applied data augmentation, TensorBoard experiment tracking, and modular evaluation pipelines to ensure reproducibility and deployment readiness",
    ]
}

EXTRA_LIST = [
    "Simon Community: Volunteered in plantation and gardening drives, contributing to community welfare and environmental care",
    "NSS (National Service Scheme): Led community drives including blood donation camps, fostering civic engagement",
    "Microsoft Space AI Event: Participated with Team CeADAR, gaining exposure to frontier AI research and industry collaboration",
    "Badminton: Represented school and college at national level; team captain demonstrating leadership and performance under pressure",
]


# ─────────────────────────────────────────────
# 2-PAGE CV
# ─────────────────────────────────────────────

def build_2page():
    doc = Document()
    set_page_margins(doc, top=0.5, bottom=0.25, left=0.4, right=0.4)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal'].font.size = Pt(FONT_SIZE)

    build_header(doc, "Data Scientist  |  AI/ML Engineer  |  Software Engineer")

    # Professional Summary
    heading(doc, "PROFESSIONAL SUMMARY", before=4, after=2)
    p_sum = doc.add_paragraph()
    p_sum.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sp(p_sum, before=0, after=2)
    r(p_sum,
        "MSc Computer Science graduate (UCD, First Class Honours) with hands-on experience in data science, "
        "machine learning, and AI engineering. Skilled in building end-to-end ML pipelines, RAG systems, and "
        "LLM-integrated applications. Proficient in Python, SQL, Java, and cloud platforms (AWS, Azure). Proven "
        "ability to deliver scalable, production-ready solutions across data engineering, NLP, and computer vision.")

    build_skills(doc, SKILLS_FULL, label_width_in=2.1, before=4)

    build_education(doc, EDUCATION_2PAGE, before=4)

    build_certifications(doc, CERTS_LINKED,
                         sec_heading="COURSES AND CERTIFICATIONS", before=4)

    heading(doc, "WORK EXPERIENCE", before=5, after=2)
    build_exp_entry(doc,
        "CeADAR – Ireland's Centre for Applied AI", "Dublin, Ireland",
        "Data Scientist Intern", "May 2025 – Present",
        CEADAR_BULLETS)
    build_exp_entry(doc,
        "LTIMindtree", "Mumbai, India",
        "Software Engineer – SQL, Java & AI Projects (Citibank Client)", "Jun 2024 – Aug 2024",
        LTIMINDTREE_BULLETS)
    build_exp_entry(doc,
        "Cognizant Technology Solutions", "Bangalore, India",
        "Programmer Analyst Trainee – Kohl's Corp, US", "Dec 2023 – May 2024",
        COGNIZANT_BULLETS)

    heading(doc, "RELEVANT PROJECTS", before=5, after=2)
    build_project(doc, PROJ_RICEAI)
    build_project(doc, PROJ_SCHOLARGENIE)
    build_project(doc, PROJ_SIAMESE)

    heading(doc, "EXTRACURRICULAR", before=5, after=2)
    for e in EXTRA_LIST:
        bul(doc, e)

    out = os.path.join(OUTPUT_DIR, "SHANTANU_CV_GENERIC_2PAGE.docx")
    doc.save(out)
    print(f"Saved: {out}")


# ─────────────────────────────────────────────
# 1-PAGE CV  (no summary, compact)
# ─────────────────────────────────────────────

def build_1page():
    doc = Document()
    set_page_margins(doc, top=0.45, bottom=0.2, left=0.4, right=0.4)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal'].font.size = Pt(FONT_SIZE)

    build_header(doc, "Data Scientist  |  AI/ML Engineer  |  Software Engineer",
                 name_size=17, tag_size=10, contact_size=8.5)

    # NO professional summary on 1-page

    build_skills(doc, SKILLS_COMPACT, label_width_in=1.65, before=3)

    build_education(doc, EDUCATION_1PAGE, before=3)

    heading(doc, "WORK EXPERIENCE", before=3, after=2)
    build_exp_entry(doc,
        "CeADAR – Ireland's Centre for Applied AI", "Dublin, Ireland",
        "Data Scientist Intern", "May 2025 – Present", CEADAR_SHORT)
    build_exp_entry(doc,
        "LTIMindtree", "Mumbai, India",
        "Software Engineer – SQL, Java & AI Projects (Citibank Client)", "Jun 2024 – Aug 2024",
        LTIMINDTREE_SHORT)
    build_exp_entry(doc,
        "Cognizant Technology Solutions", "Bangalore, India",
        "Programmer Analyst Trainee – Kohl's Corp, US", "Dec 2023 – May 2024",
        COGNIZANT_SHORT)

    heading(doc, "KEY PROJECTS", before=3, after=2)
    for proj_title, proj_date, proj_tech, proj_bullet, proj_paper_url in [
        (
            "RiceAI Expert – Sustainable Farming AI Chatbot", "CeADAR, 2025",
            "LangChain, LangGraph, Qdrant, OpenAI/Qwen, Streamlit, Neo4j, RAGAS",
            "Full-stack RAG assistant for rice farming with GraphRAG (LangGraph + Neo4j), multilingual support, and RAGAS-benchmarked retrieval quality.",
            "",
        ),
        (
            "ScholarGenie – LLM Research Discovery Agent", "March 2024",
            "Python, LangChain, GPT-4, arXiv API, Streamlit",
            "LLM-powered paper discovery tool integrating arXiv and Semantic Scholar APIs with GPT-4 summaries and interactive Streamlit interface.",
            "",
        ),
        (
            "Face Verification – Siamese CNN", "Research Paper, 2023",
            "TensorFlow, Keras, OpenCV",
            "Achieved 91.25% accuracy on 13,000 facial images using one-shot Siamese CNN; outperformed KNN baseline.",
            SIAMESE_PAPER_URL,
        ),
    ]:
        p_t = doc.add_paragraph()
        sp(p_t, before=3, after=0)
        r(p_t, proj_title, bold=True)
        # inline bold date/context
        r(p_t, "  |  ", size=FONT_SIZE)
        link_word = "Research Paper"
        if proj_paper_url and link_word in proj_date:
            idx = proj_date.index(link_word)
            if proj_date[:idx]:
                r(p_t, proj_date[:idx], bold=True)
            hyperlink(p_t, link_word, proj_paper_url)
            if proj_date[idx + len(link_word):]:
                r(p_t, proj_date[idx + len(link_word):], bold=True)
        else:
            r(p_t, proj_date, bold=True)
        p_tk = doc.add_paragraph()
        sp(p_tk, before=0, after=0)
        r(p_tk, "Technologies: " + proj_tech, italic=True)
        p_b = doc.add_paragraph()
        p_b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sp(p_b, before=0, after=1)
        r(p_b, proj_bullet)

    build_certifications(doc, CERTS_1PAGE, sec_heading="CERTIFICATIONS", before=3)

    out = os.path.join(OUTPUT_DIR, "SHANTANU_CV_GENERIC_1PAGE.docx")
    doc.save(out)
    print(f"Saved: {out}")


# ─────────────────────────────────────────────
# PDF EXPORT
# ─────────────────────────────────────────────

def export_pdfs():
    import win32com.client, shutil, tempfile
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    for f in [
        "CREATED/SHANTANU_CV_GENERIC_2PAGE.docx",
        "CREATED/SHANTANU_CV_GENERIC_1PAGE.docx",
    ]:
        docx_path = os.path.abspath(f)
        pdf_name  = os.path.splitext(os.path.basename(f))[0] + '.pdf'
        tmp_pdf   = os.path.join(tempfile.gettempdir(), pdf_name)
        final_pdf = os.path.join("CREATED", pdf_name)
        doc = word.Documents.Open(docx_path)
        doc.ExportAsFixedFormat(tmp_pdf, ExportFormat=17)
        doc.Close(False)
        if os.path.exists(final_pdf):
            os.remove(final_pdf)
        shutil.copy2(tmp_pdf, final_pdf)
        os.remove(tmp_pdf)
        print(f"PDF: {final_pdf}")
    word.Quit()


# ─────────────────────────────────────────────
# RUN
# ─────────────────────────────────────────────

if __name__ == "__main__":
    build_2page()
    build_1page()
    print("Building PDFs...")
    export_pdfs()
    print("Done.")
